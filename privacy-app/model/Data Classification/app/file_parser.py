import pandas as pd
from app.detection import calculate_confidence
from app.config import patterns



def scan_csv_with_scores(file_path, patterns):
    df = pd.read_csv(file_path)
    matches = {}
    for col in df.columns:
        for key, pattern in patterns.items():
            detected = df[col].astype(str).str.findall(pattern).sum()
            if detected:
                matches[key] = matches.get(key, []) + [
                    {"match": match, "confidence": calculate_confidence(match, key)} for match in detected
                ]
    return matches


def scan_json_with_scores(file_path, patterns):
    import json
    with open(file_path, 'r') as file:
        data = json.load(file)
    matches = {}
    for key, pattern in patterns.items():
        detected = re.findall(pattern, str(data))
        matches[key] = matches.get(key, []) + [
            {"match": match, "confidence": calculate_confidence(match, key)} for match in detected
        ]
    return matches

def scan_plain_text_with_scores(file_path, patterns):
    with open(file_path, 'r') as file:
        text = file.read()
    return detect_sensitive_data_with_scores(text, patterns)



from PyPDF2 import PdfReader
from docx import Document
import re
from app.detection import detect_sensitive_data_with_scores  # Assuming this is your detection function

def extract_text_from_pdf(file_path):
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        text += page.extract_text()
    return text

def extract_text_from_word(file_path):
    doc = Document(file_path)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    return text

def scan_pdf_with_scores(file_path, patterns):
    text = extract_text_from_pdf(file_path)
    return detect_sensitive_data_with_scores(text, patterns)

def scan_word_with_scores(file_path, patterns):
    text = extract_text_from_word(file_path)
    return detect_sensitive_data_with_scores(text, patterns)

from PyPDF2 import PdfReader
from docx import Document
from pytesseract import image_to_string
from PIL import Image
import textract
from app.detection import detect_sensitive_data_with_scores

def extract_text_from_pdf(file_path):
    try:
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        return text
    except Exception as e:
        return f"Error extracting PDF: {e}"

def extract_text_from_word(file_path):
    try:
        doc = Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text
    except Exception as e:
        return f"Error extracting Word document: {e}"

def extract_text_from_image(file_path):
    try:
        image = Image.open(file_path)
        text = image_to_string(image)
        return text
    except Exception as e:
        return f"Error extracting image text: {e}"

def extract_text_from_other(file_path):
    try:
        return textract.process(file_path).decode('utf-8')
    except Exception as e:
        return f"Error extracting text: {e}"

# file_parser.py
import json
from app.detection import detect_sensitive_data_with_scores  # Assuming detection is already defined

def scan_json(file_path):
    try:
        with open(file_path, "r") as file:
            data = json.load(file)
        text = json.dumps(data)  # Flatten nested JSON into a string
        return detect_sensitive_data_with_scores(text)
    except Exception as e:
        return {"error": f"Failed to parse JSON: {e}"}

# file_parser.py
from email import policy
from email.parser import BytesParser

def extract_text_from_email(file_path):
    try:
        with open(file_path, 'rb') as file:
            # Parse the email file
            msg = BytesParser(policy=policy.default).parse(file)

        # Extract plain text body
        body = ""
        if msg.is_multipart():
            for part in msg.iter_parts():
                if part.get_content_type() == "text/plain":
                    body += part.get_content()
        else:
            body = msg.get_content()

        return body.strip() if body else "No content found in email."
    except Exception as e:
        return f"Error extracting email content: {e}"




path = "/Users/keshavdadhich/Documents/Hackverse/uploads/test_data.csv"
from loguru import logger

logger.add("logs/{time}.log")
logger.info("Sensitive data detected in file: {}", path)




